"""Скрипт наполнения справочных данных для локальной разработки (сессия B02).
Идемпотентен: повторный запуск не создаёт дублей — каждая запись проверяется
по уникальному полю перед вставкой.

Хеширование пароля: используется passlib[argon2] напрямую (сессия B03 ещё
не создала общий app/core/security.py). Когда B03 будет сделана — с высокой
вероятностью seed-админа придётся пересоздать/перехешировать пароль через
общую функцию hash_password(), чтобы не было двух разных мест хеширования
в проекте. Технический долг, снять на B03.

Запуск (внутри контейнера backend):
    python -m app.scripts.seed_reference_data
"""
import secrets
import string

from pathlib import Path

from docx import Document
from passlib.context import CryptContext
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.database import SessionLocal
from app.models.department import Department
from app.models.document_template import DocumentTemplate, DocumentTemplateType
from app.models.equipment_type import EquipmentType
from app.models.user import User, UserRole

pwd_context = CryptContext(schemes=["argon2"], deprecated="auto")

TEMPLATES_DIR = Path(__file__).resolve().parents[2] / "storage" / "document_templates"

DEPARTMENTS = ["ИТ-отдел", "Бухгалтерия", "Регистратура", "Административно-хозяйственная часть"]

EQUIPMENT_TYPES = [
    "Компьютер",
    "Принтер",
    "Сетевое оборудование",
    "Сервер",
    "МФУ",
    "Прочее",
]


def generate_password(length: int = 16) -> str:
    alphabet = string.ascii_letters + string.digits
    return "".join(secrets.choice(alphabet) for _ in range(length))


def seed_departments(db: Session) -> dict[str, Department]:
    result = {}
    for name in DEPARTMENTS:
        existing = db.scalar(select(Department).where(Department.name == name))
        if existing:
            result[name] = existing
            continue
        dept = Department(name=name)
        db.add(dept)
        db.flush()
        result[name] = dept
        print(f"  + Department: {name}")
    return result


def seed_equipment_types(db: Session) -> None:
    for name in EQUIPMENT_TYPES:
        existing = db.scalar(select(EquipmentType).where(EquipmentType.name == name))
        if existing:
            continue
        db.add(EquipmentType(name=name))
        print(f"  + EquipmentType: {name}")


def seed_admin_user(db: Session, it_department: Department) -> None:
    existing = db.scalar(select(User).where(User.login == "admin"))
    if existing:
        print("  = Admin user already exists, skipping")
        return

    password = generate_password()
    admin = User(
        full_name="Администратор Системы",
        department_id=it_department.id,
        position="Системный администратор",
        role=UserRole.ADMIN,
        login="admin",
        password_hash=pwd_context.hash(password),
        is_active=True,
    )
    db.add(admin)
    print("  + Admin user created")
    print("  " + "=" * 50)
    print("  login:    admin")
    print(f"  password: {password}")
    print("  Сохрани этот пароль — он больше не будет показан!")
    print("  " + "=" * 50)


def create_stub_docx(path: Path, title: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph("Заглушка шаблона — реальные плейсхолдеры будут добавлены позже.")
    document.save(str(path))


def seed_document_templates(db: Session) -> None:
    templates = [
        ("Заявка на закупку", DocumentTemplateType.PURCHASE_REQUEST, "purchase_request.docx", UserRole.IT_HEAD),
        ("Акт списания", DocumentTemplateType.WRITE_OFF_ACT, "write_off_act.docx", UserRole.IT_HEAD),
        ("Наряд на работу", DocumentTemplateType.WORK_ORDER, "work_order.docx", UserRole.ENGINEER),
    ]
    for name, doc_type, filename, min_role in templates:
        existing = db.scalar(select(DocumentTemplate).where(DocumentTemplate.name == name))
        if existing:
            continue

        file_path = TEMPLATES_DIR / filename
        if not file_path.exists():
            create_stub_docx(file_path, name)

        db.add(
            DocumentTemplate(
                name=name,
                type=doc_type,
                file_path=str(file_path),
                field_schema={},
                min_approver_role=min_role,
            )
        )
        print(f"  + DocumentTemplate: {name} ({file_path})")


def main() -> None:
    db = SessionLocal()
    try:
        print("Departments:")
        departments = seed_departments(db)

        print("Equipment types:")
        seed_equipment_types(db)

        print("Admin user:")
        seed_admin_user(db, departments["ИТ-отдел"])

        print("Document templates:")
        seed_document_templates(db)

        db.commit()
        print("Готово.")
    except Exception:
        db.rollback()
        raise
    finally:
        db.close()


if __name__ == "__main__":
    main()